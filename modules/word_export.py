from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .constants import ND30_FONT_NAME, ND30_FONT_SIZE, ND30_LINE_SPACING, ND30_MARGINS_CM


def _apply_nd30_basics(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(ND30_MARGINS_CM["top"])
    section.bottom_margin = Cm(ND30_MARGINS_CM["bottom"])
    section.left_margin = Cm(ND30_MARGINS_CM["left"])
    section.right_margin = Cm(ND30_MARGINS_CM["right"])

    style = doc.styles["Normal"]
    font = style.font
    font.name = ND30_FONT_NAME
    font.size = Pt(ND30_FONT_SIZE)

    # paragraph defaults
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = ND30_LINE_SPACING


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.name = ND30_FONT_NAME
    run.font.size = Pt(ND30_FONT_SIZE)


def build_exam_docx(
    questions: List[Dict[str, Any]],
    meta: Dict[str, Any],
    template_bytes: Optional[bytes] = None,
) -> bytes:
    """Build 'Đề' docx. If template is provided, we append content after it."""
    if template_bytes:
        doc = Document(BytesIO(template_bytes))
    else:
        doc = Document()

    _apply_nd30_basics(doc)

    title = meta.get("title", "ĐỀ KIỂM TRA")
    subject = meta.get("subject", "")
    grade = meta.get("grade", "5")
    duration = meta.get("duration", "")

    _add_heading(doc, title)
    if subject:
        _add_heading(doc, f"Môn: {subject}  |  Lớp: {grade}")
    if duration:
        p = doc.add_paragraph(f"Thời gian làm bài: {duration}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(" ")

    # Group by qtype
    idx = 1
    for q in questions:
        qtype = q.get("qtype", "")
        points = q.get("points", "")
        content = q.get("content", {}) or {}

        stem = content.get("stem") or content.get("prompt") or ""
        p = doc.add_paragraph(f"Câu {idx}. ({points} điểm) {stem}")

        if qtype == "Trắc nghiệm nhiều lựa chọn":
            opts = (content.get("options") or {})
            for k in ["A", "B", "C", "D"]:
                doc.add_paragraph(f"{k}. {opts.get(k, '')}")
        elif qtype == "Đúng/Sai":
            for j, st in enumerate(content.get("statements") or [], start=1):
                doc.add_paragraph(f"{j}) {st.get('text','')}")
        elif qtype == "Nối cột":
            left = content.get("left") or []
            right = content.get("right") or []
            table = doc.add_table(rows=max(len(left), len(right)) + 1, cols=2)
            table.cell(0, 0).text = "Cột A"
            table.cell(0, 1).text = "Cột B"
            for r in range(max(len(left), len(right))):
                table.cell(r + 1, 0).text = str(left[r]) if r < len(left) else ""
                table.cell(r + 1, 1).text = str(right[r]) if r < len(right) else ""
        elif qtype == "Điền khuyết":
            doc.add_paragraph(content.get("text", ""))

        doc.add_paragraph(" ")
        idx += 1

    # Answer key
    doc.add_page_break()
    _add_heading(doc, "ĐÁP ÁN")

    for i, q in enumerate(questions, start=1):
        qtype = q.get("qtype", "")
        content = q.get("content", {}) or {}
        if qtype == "Trắc nghiệm nhiều lựa chọn":
            doc.add_paragraph(f"Câu {i}: {content.get('answer','')}")
        elif qtype == "Đúng/Sai":
            ans = []
            for st in content.get("statements") or []:
                a = st.get("answer")
                if a in [True, "Đ", "Đúng", "True"]:
                    ans.append("Đ")
                elif a in [False, "S", "Sai", "False"]:
                    ans.append("S")
                else:
                    ans.append("?")
            doc.add_paragraph(f"Câu {i}: {' '.join(ans)}")
        elif qtype == "Nối cột":
            doc.add_paragraph(f"Câu {i}: {content.get('mapping', {})}")
        elif qtype == "Điền khuyết":
            doc.add_paragraph(f"Câu {i}: {content.get('answer','')}")
        elif qtype == "Tự luận":
            doc.add_paragraph(f"Câu {i}: (GV chấm theo rubric)")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_blueprint_docx(
    matrix_records: List[Dict[str, Any]],
    meta: Dict[str, Any],
    template_bytes: Optional[bytes] = None,
) -> bytes:
    """Build 'Bảng đặc tả' docx (simple but consistent)."""
    if template_bytes:
        doc = Document(BytesIO(template_bytes))
    else:
        doc = Document()

    _apply_nd30_basics(doc)

    title = meta.get("blueprint_title", "BẢNG ĐẶC TẢ ĐỀ")
    _add_heading(doc, title)

    cols = [
        "TT",
        "Chủ đề",
        "Bài/Nội dung",
        "YCCĐ",
        "Số tiết",
        "Tỉ lệ (%)",
        "Số điểm cần đạt",
        "Tổng số câu/ý",
    ]

    table = doc.add_table(rows=1, cols=len(cols))
    for j, c in enumerate(cols):
        table.cell(0, j).text = c

    for rec in matrix_records:
        row = table.add_row().cells
        for j, c in enumerate(cols):
            row[j].text = str(rec.get(c, ""))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
